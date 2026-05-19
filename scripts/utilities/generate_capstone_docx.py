from pathlib import Path
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

src = Path(r"C:\Final\DesignXcel01\docs\CAPSTONE_PAPER_ALIGNED_TO_SYSTEM.md")
out = Path(r"C:\Final\DesignXcel01\docs\CAPSTONE_PAPER_ALIGNED_TO_SYSTEM.docx")

lines = src.read_text(encoding="utf-8").splitlines()

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

def add_runs_with_bold(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

for raw in lines:
    line = raw.rstrip()
    if not line:
        doc.add_paragraph("")
        continue

    if line.startswith("# "):
        p = doc.add_heading(line[2:].strip(), level=1)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        continue

    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
        continue

    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
        continue

    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_runs_with_bold(p, line[2:].strip())
        continue

    p = doc.add_paragraph()
    add_runs_with_bold(p, line)

while doc.paragraphs and not doc.paragraphs[-1].text.strip():
    p = doc.paragraphs[-1]._element
    p.getparent().remove(p)

doc.save(out)
print(out)
